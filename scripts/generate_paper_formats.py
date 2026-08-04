"""
Generate publication-ready DOCX (.docx) and LaTeX (.tex) formats
for the ScopeGuard research paper.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    """Set background color for a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def generate_docx():
    """Build formatted Word document (IEEE/ACM paper format)."""
    doc = docx.Document()

    # Set page margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ScopeGuard: Identity-Scoped Execution Control and Dynamic Permission Authorization for AI Agents in Cloud Infrastructure\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 51, 102)

    # Authors
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_authors = p_authors.add_run("Binod Prasad Joshi\nDepartment of Computer Science & Cloud Infrastructure Security Research Group\n\n")
    run_authors.font.name = 'Times New Roman'
    run_authors.font.size = Pt(11)
    run_authors.font.italic = True

    # Abstract Box
    p_abs_hdr = doc.add_paragraph()
    run_abs_hdr = p_abs_hdr.add_run("Abstract")
    run_abs_hdr.font.name = 'Times New Roman'
    run_abs_hdr.font.size = Pt(12)
    run_abs_hdr.font.bold = True

    p_abs = doc.add_paragraph()
    run_abs = p_abs.add_run(
        "The rapid integration of Large Language Model (LLM)-powered autonomous agents into cloud management workflows promises to automate complex infrastructure-as-code (IaC) provisioning, resource scaling, and operational monitoring. However, deploying off-the-shelf AI agents with static cloud credentials introduces severe security risks: agents granted broad Identity and Access Management (IAM) permissions can exhibit unintended scope creep, succumb to prompt injection attacks, or execute destructive actions without deterministic identity attribution. Traditional access control models—such as static Role-Based Access Control (RBAC) and declarative Policy-as-Code engines (e.g., OPA, AWS Cedar)—fail to dynamically infer permission boundaries from natural language task objectives prior to execution.\n\n"
        "We present ScopeGuard, an asynchronous, identity-scoped proxy framework designed for real-time permission authorization and execution control of AI agents in cloud infrastructure. ScopeGuard intercepts outbound agent tool calls, attaches cryptographically verifiable identity metadata (Agent ID, Session UUID, declared task, and reasoning chain), and dynamically compiles a task-aware permission scope using a rule-based action taxonomy. Intercepted API calls are evaluated against this computed scope under a fail-closed policy engine supporting three ablation operational modes (passthrough, tagging_only, and full).\n\n"
        "We empirically evaluate ScopeGuard across a benchmark suite of 200 infrastructure management tasks (100 legitimate operations and 100 adversarial/off-scope attacks across 10 security threat categories). In full enforcement mode, ScopeGuard achieves an overall accuracy of 89.0%, a scoping precision of 97.8%, an adversarial block rate of 88.0%, and a false positive rate of 2.0% on legitimate tasks, while introducing a minimal average latency overhead of 108.9 ms—well within the 200 ms real-time inference budget. We validate live execution against real Amazon Web Services (AWS) infrastructure endpoints including Amazon RDS, EC2, S3, and CloudWatch."
    )
    run_abs.font.name = 'Times New Roman'
    run_abs.font.size = Pt(10)
    run_abs.font.italic = True

    # Keywords
    p_kw = doc.add_paragraph()
    run_kw_title = p_kw.add_run("Keywords—")
    run_kw_title.font.bold = True
    run_kw_title.font.size = Pt(10)
    run_kw_val = p_kw.add_run("AI Agent Security, Infrastructure Access Control, Cloud Governance, Prompt Injection Defense, Least Privilege, Real-time Authorization Proxy, AWS Identity Scoping.\n\n")
    run_kw_val.font.size = Pt(10)

    # Section I
    h1 = doc.add_heading("I. INTRODUCTION", level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)

    p1 = doc.add_paragraph(
        "Cloud infrastructure environments are the foundation of modern digital applications, relying on complex control planes to manage compute instances, relational databases, object storage, and network topologies. Managing these resources requires executing multi-step infrastructure operations via Cloud APIs (e.g., AWS SDK boto3), Infrastructure-as-Code (IaC) CLI tools (e.g., Terraform), and CI/CD pipelines (e.g., GitHub Actions)."
    )

    # Section II
    doc.add_heading("II. BACKGROUND AND MOTIVATION", level=1)
    doc.add_paragraph("Traditional cloud Security models grant static permissions to execution roles. When autonomous agents operate dynamically, over-provisioning leads to prompt injection exposure, while under-provisioning leads to operational failure.")

    # Section III
    doc.add_heading("III. SCOPEGUARD SYSTEM ARCHITECTURE", level=1)
    doc.add_paragraph("ScopeGuard functions as an intercepting reverse proxy between AI agent tool-use interfaces and cloud infrastructure backend APIs. The proxy pipeline contains five primary modules: Identity Tagger, Task-to-Scope Classifier, Policy Engine, SQLite Audit Logger, and Backend Executor.")

    # Table 1: Ablation Summary
    doc.add_heading("IV. EMPIRICAL BENCHMARK EVALUATION", level=1)
    doc.add_paragraph("Table 1 summarizes the empirical ablation results across 600 tool call executions:")

    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["Metric / Dimension", "Passthrough", "Tagging Only", "ScopeGuard (Full)"]
    for i, h_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.paragraphs[0].text = h_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ["Total Calls Evaluated", "200", "200", "200"],
        ["Classification Accuracy", "50.0%", "50.0%", "89.0%"],
        ["Scoping Precision", "0.00", "0.00", "0.978 (97.8%)"],
        ["Adversarial Block Rate", "0.0%", "0.0%", "88.0%"],
        ["False Positive Rate", "0.0%", "0.0%", "2.0%"],
        ["F1 Score", "0.00", "0.00", "0.926"],
        ["Avg Latency Overhead", "129.2 ms", "111.2 ms", "108.9 ms"],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        bg = "F4F6F9" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].text = cell_value
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_background(cell, bg)

    doc.add_paragraph("\n")

    # Section V
    doc.add_heading("V. CONCLUSION AND FUTURE WORK", level=1)
    doc.add_paragraph("ScopeGuard provides an identity-scoped execution control proxy for AI agents operating in cloud infrastructure. In empirical evaluation, it achieved an 89.0% accuracy and 88.0% adversarial block rate with sub-110ms latency overhead.")

    # Save
    doc.save("paper/ScopeGuard_Research_Paper.docx")
    print("Successfully generated paper/ScopeGuard_Research_Paper.docx")


def generate_tex():
    """Build LaTeX manuscript file (IEEEtran format)."""
    tex_content = r"""\documentclass[conference]{IEEEtran}
\usepackage{cite}
\usepackage{amsmath,amssymb,amsfonts}
\usepackage{algorithmic}
\usepackage{graphicx}
\usepackage{textcomp}
\usepackage{xcolor}
\usepackage{booktabs}

\begin{document}

\title{ScopeGuard: Identity-Scoped Execution Control and Dynamic Permission Authorization for AI Agents in Cloud Infrastructure}

\author{\IEEEauthorblockN{Binod Prasad Joshi}
\IEEEauthorblockA{\textit{Department of Computer Science \& Cloud Infrastructure Security} \\
\textit{Research Group}\\
Email: research@scopeguard.io}
}

\maketitle

\begin{abstract}
The rapid integration of Large Language Model (LLM)-powered autonomous agents into cloud management workflows promises to automate complex infrastructure-as-code (IaC) provisioning, resource scaling, and operational monitoring. However, deploying off-the-shelf AI agents with static cloud credentials introduces severe security risks: agents granted broad Identity and Access Management (IAM) permissions can exhibit unintended scope creep, succumb to prompt injection attacks, or execute destructive actions without deterministic identity attribution. We present ScopeGuard, an asynchronous, identity-scoped proxy framework designed for real-time permission authorization and execution control of AI agents in cloud infrastructure. In full enforcement mode, ScopeGuard achieves an overall accuracy of 89.0\%, a scoping precision of 97.8\%, an adversarial block rate of 88.0\%, and a false positive rate of 2.0\%, with a minimal average latency overhead of 108.9 ms.
\end{abstract}

\begin{IEEEkeywords}
AI Agent Security, Infrastructure Access Control, Cloud Governance, Prompt Injection Defense, Real-time Authorization Proxy.
\end{IEEEkeywords}

\section{Introduction}
Cloud infrastructure environments are the foundation of modern digital applications, relying on complex control planes to manage compute instances, relational databases, object storage, and network topologies. Managing these resources requires executing multi-step infrastructure operations via Cloud APIs (e.g., AWS SDK boto3) and Infrastructure-as-Code CLI tools (e.g., Terraform).

\section{System Architecture}
ScopeGuard functions as an intercepting reverse proxy between an AI Agent execution environment and cloud infrastructure backends.

\section{Experimental Evaluation}
\begin{table}[htbp]
\caption{ScopeGuard Ablation Evaluation Metrics (600 Tool Calls)}
\begin{center}
\begin{tabular}{lrrr}
\toprule
\textbf{Metric} & \textbf{Passthrough} & \textbf{Tagging Only} & \textbf{ScopeGuard Full} \\
\midrule
Accuracy & 50.0\% & 50.0\% & \textbf{89.0\%} \\
Precision & 0.00 & 0.00 & \textbf{0.978} \\
Recall (Block Rate) & 0.00 & 0.00 & \textbf{0.880} \\
FPR & 0.0\% & 0.0\% & \textbf{2.0\%} \\
Latency Overhead & 129.2 ms & 111.2 ms & \textbf{108.9 ms} \\
\bottomrule
\end{tabular}
\end{center}
\end{table}

\section{Conclusion}
ScopeGuard bridges the critical security gap between autonomous agent capabilities and cloud governance, achieving high security efficacy with sub-110ms execution overhead.

\end{document}
"""
    with open("paper/scopeguard_research_paper.tex", "w") as f:
        f.write(tex_content)
    print("Successfully generated paper/scopeguard_research_paper.tex")


if __name__ == "__main__":
    generate_docx()
    generate_tex()
